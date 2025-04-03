import fitz  # PyMuPDF
import pdfplumber
import pinecone
import re
from PIL import Image
import io
from pinecone import Pinecone, ServerlessSpec
import base64  # For encoding/decoding image data
from io import BytesIO
from docx import Document
import os
import docx
import uuid
from .model import generate_embedding

class DocumentProcessor:
    def __init__(self, pinecone_api_key=None, pinecone_index_name=None):
        """Initialize the document processor with Pinecone credentials."""
        # Initialize Pinecone if credentials are provided
        if pinecone_api_key and pinecone_index_name:
            self.pc = Pinecone(api_key=pinecone_api_key)
            self.index = self.pc.Index(pinecone_index_name)
        else:
            self.pc = None
            self.index = None

        # Initialize text splitting parameters
        self.chunk_size = 1000
        self.chunk_overlap = 200

    # def process_documents(self, file_path):
    #     """Process a document file and store chunks in Pinecone."""
    #     if not os.path.exists(file_path):
    #         raise FileNotFoundError(f"Document file not found: {file_path}")
            
    #     # Extract text from document
    #     text = self.extract_text(file_path)
        
    #     # Split into chunks
    #     chunks = self.split_text(text)
    #     print(f"Split document into {len(chunks)} chunks")
        
    #     # Create vectors and store in Pinecone
    #     if self.index:
    #         file_name = os.path.basename(file_path)
    #         vectors = self.create_vectors(chunks, file_name)
            
    #         # Upsert in batches
    #         batch_size = 100
    #         for i in range(0, len(vectors), batch_size):
    #             batch = vectors[i:i+batch_size]
    #             self.index.upsert(vectors=batch)
                
    #         print(f"Uploaded {len(vectors)} vectors to Pinecone")
            
    #     return chunks
        
    # def extract_text(self, file_path):
    #     """Extract text from a DOCX file."""
    #     try:
    #         doc = docx.Document(file_path)
    #         full_text = []
            
    #         for para in doc.paragraphs:
    #             full_text.append(para.text)
                
    #         return '\n'.join(full_text)
    #     except Exception as e:
    #         raise Exception(f"Error extracting text from document: {str(e)}")
            
    def split_text(self, text):
        """Split text into chunks of specified size."""
        chunks = []
        current_chunk = ""
        current_size = 0
        
        # Split by paragraphs
        paragraphs = text.split('\n')
        
        for para in paragraphs:
            # Skip empty paragraphs
            if not para.strip():
                continue
                
            # If adding this paragraph would exceed chunk size, save current chunk
            if current_size + len(para) > self.chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                
                # Start new chunk with overlap
                words = current_chunk.split()
                overlap_words = words[-self.chunk_overlap:] if len(words) > self.chunk_overlap else words
                current_chunk = ' '.join(overlap_words) + '\n' + para
                current_size = len(current_chunk)
            else:
                # Add to current chunk
                current_chunk += '\n' + para if current_chunk else para
                current_size += len(para)
                
        # Add the last chunk if not empty
        if current_chunk:
            chunks.append(current_chunk.strip())
            
        return chunks
        
    # def create_vectors(self, chunks, filename):
    #     """Create vectors from text chunks using generate_embedding."""
    #     vectors = []
        
    #     for i, chunk in enumerate(chunks):
    #         # Generate a unique ID
    #         chunk_id = f"{uuid.uuid4()}"
            
    #         # Create metadata
    #         metadata = {
    #             "text": chunk,
    #             "type": "text",
    #             "source": filename,
    #             "chunk_index": i
    #         }
            
    #         # Generate embedding using the model.py function
    #         embedding = generate_embedding(chunk)
            
    #         # Create vector
    #         vector = {
    #             "id": chunk_id,
    #             "values": embedding,
    #             "metadata": metadata
    #         }
            
    #         vectors.append(vector)
            
    #     return vectors

    def extract_content(self, doc_path):
        """Extract text and tables from various document types"""
        text = ""
        tables = []

        if doc_path.endswith('.docx'):
            doc = Document(doc_path)
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            for para in doc.paragraphs:
                text += para.text + "\n"

        elif doc_path.endswith('.txt'):
            with open(doc_path, 'r', encoding='utf-8') as file:
                text = file.read()

        # Add more document types as needed

        return text, tables

    def process_documents(self, doc_path):
        """Full document processing pipeline"""
        text, tables = self.extract_content(doc_path)

        # Prepare table chunks
        table_chunks = []
        for i, table in enumerate(tables):
            if table:
                table_text = ""
                for row in table:
                    if row:
                        cleaned_row = [str(cell) if cell is not None else "" for cell in row]
                        table_text += "|".join(cleaned_row) + "\n"
                table_chunks.append({"id": f"doc_table_{i}", "text": table_text, "type": "table"})

        # Prepare text chunks
        text_chunks = self.split_text(text)
        text_chunks = [{
            "id": f"doc_text_{i}", 
            "text": chunk, 
            "type": "text"} for i, chunk in enumerate(text_chunks)]

        # Combine all chunks
        all_chunks = text_chunks + table_chunks

        # Debug: Print the number of chunks
        print(f"Number of chunks: {len(all_chunks)} (Text: {len(text_chunks)}, Tables: {len(table_chunks)})")

        # Generate embeddings and prepare vectors for Pinecone
        vectors = []
        for chunk in all_chunks:
            # Generate embedding for each chunk
            embedding = generate_embedding(chunk["text"])
            
            vector = {
                "id": chunk["id"],
                "values": embedding,
                "metadata": {
                    "text": chunk["text"],
                    "type": chunk["type"]
                }
            }
            vectors.append(vector)

        # Debug: Print vector details
        print(f"Number of vectors to be uploaded: {len(vectors)}")
        print(f"Example vector metadata: {vectors[0]['metadata']}")

        # Upsert vectors to Pinecone
        for i in range(0, len(vectors), 50):
            batch = vectors[i:i+50]
            print(f"Uploading batch {i//100 + 1} with {len(batch)} vectors")
            try:
                response = self.index.upsert(vectors=batch)
                print(f"Batch {i//100 + 1} uploaded successfully")
            except Exception as e:
                print(f"Error upserting batch {i//100 + 1}: {str(e)}")

        return all_chunks